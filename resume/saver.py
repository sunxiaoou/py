#! /usr/local/bin/python3

import os
import datetime
import docx

from resume import Keys


class Saver:

    @staticmethod
    def add_reason(doc):
        head = doc.add_heading('推荐理由', 1)
        head.runs[0].bold = True
        doc.add_paragraph('', 'List Number')
        doc.add_paragraph('', 'List Number')
        doc.add_paragraph('', 'List Number')

    @staticmethod
    def add_person(doc, resume):
        head = doc.add_heading('候选人基本信息', 1)
        head.runs[0].bold = True

        para = doc.add_paragraph('姓名：{:<40} 性别： {}'.format(resume.get(Keys.name), resume.get(Keys.gender)))
        para.runs[0].add_break()

        years = age = -1
        today = datetime.datetime.today()
        year = resume.get(Keys.year)
        if year is not None:
            years = today.year - year
        birth = resume.get(Keys.birth)
        if birth is not None:
            age = today.year - birth.year - ((today.month, today.day) < (birth.month, birth.day))
        para.add_run('工作年限：{:<37d} 年龄：{:2d}'.format(years, age))
        para.runs[1].add_break()

        degree = spot = None
        education = resume.get(Keys.education)
        if education is not None:
            degree = ['大专', '本科', '硕士', 'MBA', 'EMBA', '博士',  '博士后'][education - 1]
            enrollment = resume.get(Keys.enrollment)
            if enrollment is not None:
                degree += ' ({})'.format(enrollment)
        spots = resume.get(Keys.spots)
        if spots and spots is not None:
            spot = spots[0]
        para.add_run('学历：{:<40} 目前工作地：{}'.format(degree, spot))

    @staticmethod
    def add_one_item(para, name, value):
        if value:
            run = para.add_run('{}:'.format(name))
            run.bold = True
            run.add_break()
            run = para.add_run(value)
            run.add_break()

    @staticmethod
    def add_experiences(doc, resume):
        head = doc.add_heading(Keys.experiences, 1)
        head.runs[0].bold = True

        s = resume.get(Keys.experiences)
        if s is None:
            return
        experiences = eval(s)
        for experience in experiences:
            date1 = experience.get(Keys.start_date)
            if date1:
                date1 = date1.strftime('%Y/%m')
            date2 = experience.get(Keys.end_date)
            if date2:
                date2 = date2.strftime('%Y/%m')
            company = experience.get(Keys.company)
            job = experience.get(Keys.job)
            job_desc = experience.get(Keys.job_desc)

            para = doc.add_paragraph('{}-{}: {}'.format(date1, date2, company))
            para.runs[0].bold = True
            para.runs[0].add_break()
            Saver.add_one_item(para, Keys.job, job)
            Saver.add_one_item(para, Keys.job_desc, job_desc)

    @staticmethod
    def add_projects(doc, resume):
        head = doc.add_heading(Keys.projects, 1)
        head.runs[0].bold = True

        s = resume.get(Keys.projects)
        if s is None:
            return
        projects = eval(s)
        for project in projects:
            date1 = project.get(Keys.start_date)
            if date1:
                date1 = date1.strftime('%Y/%m')
            date2 = project.get(Keys.end_date)
            if date2:
                date2 = date2.strftime('%Y/%m')
            name = project.get(Keys.project)
            description = project.get(Keys.project_desc)
            duty = project.get(Keys.duty)

            para = doc.add_paragraph('{}-{}: {}'.format(date1, date2, name))
            para.runs[0].bold = True
            para.runs[0].add_break()
            Saver.add_one_item(para, Keys.project_desc, description)
            Saver.add_one_item(para, Keys.duty, duty)

    @staticmethod
    def add_educations(doc, resume):
        head = doc.add_heading(Keys.educations, 1)
        head.runs[0].bold = True

        s = resume.get(Keys.educations)
        if s is None:
            return
        educations = eval(s)
        for education in educations:
            date1 = education.get(Keys.start_date)
            if date1:
                date1 = date1.strftime('%Y/%m')
            date2 = education.get(Keys.end_date)
            if date2:
                date2 = date2.strftime('%Y/%m')
            school = education.get(Keys.school)
            major = education.get(Keys.major)
            degree = education.get(Keys.degree)
            if degree:
                degree = ['大专', '本科', '硕士', 'MBA', 'EMBA', '博士',  '博士后'][degree - 1]
            para = doc.add_paragraph('{}-{}: {}'.format(date1, date2, school))
            para.runs[0].bold = True
            para.runs[0].add_break()
            run = para.add_run('{}:  {}'.format(Keys.major, major))
            run.add_break()
            para.add_run('{}:  {}'.format(Keys.degree, degree))

    @staticmethod
    def add_languages(doc, resume):
        head = doc.add_heading(Keys.languages, 1)
        head.runs[0].bold = True

        s = resume.get(Keys.languages)
        if s is None:
            return
        doc.add_paragraph(s)

    @staticmethod
    def to_doc(resume, file):
        doc = docx.Document()
        doc.add_paragraph('推荐简历', 'Title')

        doc.add_paragraph('应聘岗位：')

        Saver.add_reason(doc)
        Saver.add_person(doc, resume)
        Saver.add_experiences(doc, resume)
        Saver.add_projects(doc, resume)
        Saver.add_educations(doc, resume)
        Saver.add_languages(doc, resume)

        doc.save(file)


def main():
    file = open('jl_0043079_陈磊.txt')
    # file = open('jxw_0051701_龙浩.txt')
    s = file.read()
    resume = eval(s)
    html = resume.get(Keys.file)
    Saver.to_doc(resume, os.path.splitext(html)[0] + '.docx')

if __name__ == "__main__":
    main()
