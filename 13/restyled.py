#! /usr/bin/python3

import docx

doc = docx.Document('demo.docx')
print(doc.paragraphs[0].text)
# 'Document Title'
print(doc.paragraphs[0].style.name)
# 'Title'
doc.paragraphs[0].style = 'Normal'

print(doc.paragraphs[1].text)
# # 'A plain paragraph with some bold and some italic'
doc.paragraphs[1].runs[0].style.name = 'QuoteChar'
doc.paragraphs[1].runs[2].underline = True
doc.paragraphs[1].runs[4].underline = True
doc.save('restyled.docx')
