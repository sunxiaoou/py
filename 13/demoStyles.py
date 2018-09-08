#! /usr/bin/python3

import docx

doc = docx.Document('demo.docx')

# for p in doc.paragraphs:
for i in range(len(doc.paragraphs)):
    p = doc.paragraphs[i]
    print('{:02d}: {} ({})'.format(i, p.text, p.style.name))
    for r in p.runs:
        attributes = []
        if r.bold:
            attributes.append('bold')
        if r.italic:
            attributes.append('italic')
        if r.underline:
            attributes.append('underline')
        print('    {} ({}){}'.format(r.text, r.style.name, str(attributes)))
    print()
